import os
import fitz  # PyMuPDF
import docx  # python-docx
from flask import Flask, request, jsonify, redirect, url_for, render_template
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.utils import secure_filename
from modules.genai_engine import GenAIEngine
from models import db, User, Document # Import DB and Models
from utils import set_password, check_password

# --- CONFIGURATION ---
# The database file will be named vault.db and stored in the backend folder
DB_NAME = 'vault.db'
UPLOAD_FOLDER = 'uploads'

# 1. Initialize Flask App
# static_url_path='' tells Flask to serve style.css from the root, not /static
app = Flask(__name__, 
            static_folder='../frontend', 
            template_folder='../frontend',
            static_url_path='')
CORS(app)

# --- Security & Database Setup ---
app.config['SECRET_KEY'] = 'your_super_secret_key_for_sessions' # CRUCIAL: Change this in production!
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{DB_NAME}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Initialize DB and Login Manager
db.init_app(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login' # Where to redirect if user isn't logged in

# Initialize GenAI Engine
engine = GenAIEngine()

# --- Database Setup (Happens once) ---
with app.app_context():
    # Create the database file if it doesn't exist
    db.create_all()

# --- HELPER FUNCTIONS ---
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = "".join([page.get_text() for page in doc])
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def process_file(filepath):
    ext = filepath.split('.')[-1].lower()
    if ext == 'pdf': return extract_text_from_pdf(filepath)
    elif ext == 'docx': return extract_text_from_docx(filepath)
    elif ext == 'txt': 
        with open(filepath, 'r', encoding='utf-8') as f: return f.read()
    return None

# --- Flask-Login User Loader ---
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# --- WEB ROUTES (HTML PAGES) ---
# These routes serve the HTML files to the user
@app.route('/')
def index_redirect():
    # If not logged in, show login page. If logged in, show the app.
    if current_user.is_authenticated:
        return app.send_static_file('index.html')
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        # Authentication Logic
        data = request.json
        user = User.query.filter_by(username=data['username']).first()
        
        if user and check_password(user.password_hash, data['password']):
            login_user(user)
            # Fetch user's documents for the dashboard
            documents = Document.query.filter_by(user_id=user.id).order_by(Document.upload_date.desc()).limit(10).all()
            doc_list = [{"id": d.id, "filename": d.filename} for d in documents]
            return jsonify({'success': True, 'documents': doc_list})
        
        return jsonify({'success': False, 'message': 'Invalid username or password.'}), 401
    
    return app.send_static_file('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        # Registration Logic
        data = request.json
        
        if User.query.filter_by(username=data['username']).first():
            return jsonify({'success': False, 'message': 'Username already exists.'}), 409
        
        new_user = User(
            username=data['username'],
            password_hash=set_password(data['password'])
        )
        
        try:
            db.session.add(new_user)
            db.session.commit()
            return jsonify({'success': True, 'message': 'Registration successful!'})
        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500
    
    return app.send_static_file('register.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))


# --- API ENDPOINTS (Protected) ---
@app.route('/api/upload', methods=['POST'])
@login_required
def upload_file():
    # NOTE: This logic is nearly the same, but now it records the upload in the DB
    if 'file' not in request.files: return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    
    try:
        extracted_text = process_file(filepath)
        if not extracted_text or len(extracted_text) < 100:
            raise ValueError("Could not extract enough text from file.")
            
        # 1. Index Memory
        collection_id = f"user_{current_user.id}_doc_{filename}_{datetime.now().timestamp()}"
        status = engine.build_memory_index(extracted_text, collection_id=collection_id)
        
        # 2. Record in DB
        new_document = Document(
            filename=filename,
            chroma_collection_id=collection_id,
            user_id=current_user.id
        )
        db.session.add(new_document)
        db.session.commit()
        
        # 3. Generate content
        summary = engine.generate_summary(extracted_text)
        quiz = engine.generate_quiz(extracted_text, count=10)
        videos = engine.get_youtube_recommendations(extracted_text)
        
        return jsonify({
            "message": status,
            "summary": summary,
            "quiz": quiz,
            "videos": videos,
            "raw_text": extracted_text, # Passed to frontend for quiz re-generation
            "document_id": new_document.id
        })
        
    except Exception as e:
        db.session.rollback() # Rollback if database entry failed
        return jsonify({"error": f"Processing Error: {str(e)}", "message": "Failed to process document."}), 500
        
    finally:
        if os.path.exists(filepath): os.remove(filepath)

@app.route('/api/chat', methods=['POST'])
@login_required
def chat():
    # Chat logic remains the same, but protected by login_required
    data = request.json
    question = data.get('question', '')
    if not question: return jsonify({"error": "No question"}), 400
    answer = engine.chat_with_memory(question)
    return jsonify({"answer": answer})

@app.route('/api/quiz-more', methods=['POST'])
@login_required
def generate_more_quiz():
    # Quiz logic remains the same, but protected by login_required
    data = request.json
    text = data.get('text', '')
    difficulty = data.get('difficulty', 'Medium')
    count = data.get('count', 10) 
    if count > 10: count = 10
        
    try:
        quiz = engine.generate_quiz(text, difficulty=difficulty, count=count)
        return jsonify({"quiz": quiz})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/documents', methods=['GET'])
@login_required
def get_user_documents():
    # Returns the user's document history
    documents = Document.query.filter_by(user_id=current_user.id).order_by(Document.upload_date.desc()).limit(10).all()
    doc_list = [{"id": d.id, "filename": d.filename} for d in documents]
    return jsonify({"documents": doc_list})


if __name__ == '__main__':
    # Create the uploads folder outside the application root
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(debug=True, port=5000)